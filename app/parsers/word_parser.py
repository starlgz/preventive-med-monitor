import docx
import re
from typing import List, Dict, Any, Optional
from app.core.logger import logger
from app.parsers.column_mapper import ColumnMapper

class WordJobParser:
    """Word 文档 (.docx) 岗位表专用解析器"""

    @classmethod
    def parse_file(cls, file_path: str, default_unit_name: Optional[str] = None) -> List[Dict[str, Any]]:
        jobs = []
        try:
            doc = docx.Document(file_path)
            for table in doc.tables:
                table_jobs = cls._parse_table(table, default_unit_name)
                jobs.extend(table_jobs)
            # 如果表格中未解析到岗位，尝试从段落文本列表中提取
            if not jobs:
                jobs = cls._parse_paragraphs(doc, default_unit_name)
        except Exception as e:
            logger.error(f"Error parsing word docx {file_path}: {e}")
        return jobs

    @classmethod
    def _parse_paragraphs(cls, doc: docx.Document, default_unit_name: Optional[str] = None) -> List[Dict[str, Any]]:
        """当 Word 附件使用段落列表而非表格时进行文本规则提取"""
        jobs = []
        current_unit = default_unit_name or "未指定招聘单位"
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            # 过滤纯标题行，如“一、招聘岗位及要求”、“附件：岗位表”
            if re.match(r"^[一二三四五六七八九十]+[、\.\s]|(?:附[件表]|通知|公告|说明|要求)$", text):
                if not re.search(r"人|专业|学历|周岁", text):
                    continue

            # 检测是否包含岗位特征（例如：“1. 流行病学科流调员：招聘2人，要求预防医学专业，本科及以上学历”）
            job_pattern = re.search(r"^(?:\d+[\.、\s]+)?([^\n:：]{2,20}?(?:医师|人员|岗位|岗|职位|技师|科|室|员))[:：\s]+(.*)", text)
            if job_pattern:
                job_title = job_pattern.group(1).strip()
                # 再次剔除可能残留的前缀序号如 "1. " 或 "(1)"
                job_title = re.sub(r"^[\d一二三四五六七八九十]+[\.、\s\)\(]+", "", job_title).strip()
                desc = job_pattern.group(2).strip()

                # 排除例如“一、招聘岗位”或纯词语
                if job_title in ["招聘岗位", "岗位要求", "报考条件", "基本条件", "应聘人员"]:
                    continue

                # 提取人数
                hc_match = re.search(r"(?:招[聘录]|计划|共)?(\d+)\s*人", desc)
                headcount = int(hc_match.group(1)) if hc_match else 1

                # 提取学历
                edu_match = re.search(r"(博士研究生|硕士研究生|研究生|博士|硕士|本科|大学本科|大专|专科)(?:及以上)?", desc)
                education = edu_match.group(0) if edu_match else ""

                # 提取专业
                major_match = re.search(r"(?:专业|要求|所学专业)[:：\s]*([^，,；;。]+)", desc)
                major_raw = major_match.group(1).strip() if major_match else desc

                jobs.append({
                    "unit_name": current_unit,
                    "job_name": job_title,
                    "job_code": "",
                    "headcount": headcount,
                    "education": education,
                    "major_raw": major_raw,
                    "other_requirements": desc
                })
        return jobs

    @classmethod
    def _parse_table(cls, table, default_unit_name: Optional[str]) -> List[Dict[str, Any]]:
        rows = table.rows
        if not rows:
            return []

        header_row_idx = -1
        col_map = {}
        for idx, row in enumerate(rows):
            row_vals = [cell.text.strip() for cell in row.cells]
            mapped = ColumnMapper.map_columns(row_vals)
            if "job_name" in mapped or "major" in mapped:
                header_row_idx = idx
                col_map = mapped
                break

        if header_row_idx == -1:
            return []

        jobs = []
        last_unit = default_unit_name or ""
        import re
        for r_idx in range(header_row_idx + 1, len(rows)):
            row = rows[r_idx]
            row_vals = [cell.text.strip() for cell in row.cells]
            if not any(row_vals):
                continue

            unit = row_vals[col_map["unit_name"]] if "unit_name" in col_map and col_map["unit_name"] < len(row_vals) else ""
            job_name = row_vals[col_map["job_name"]] if "job_name" in col_map and col_map["job_name"] < len(row_vals) else ""
            major_raw = row_vals[col_map["major"]] if "major" in col_map and col_map["major"] < len(row_vals) else ""
            job_code = row_vals[col_map["job_code"]] if "job_code" in col_map and col_map["job_code"] < len(row_vals) else ""
            headcount_raw = row_vals[col_map["headcount"]] if "headcount" in col_map and col_map["headcount"] < len(row_vals) else "1"
            education = row_vals[col_map["education"]] if "education" in col_map and col_map["education"] < len(row_vals) else ""
            degree = row_vals[col_map["degree"]] if "degree" in col_map and col_map["degree"] < len(row_vals) else ""
            age = row_vals[col_map["age"]] if "age" in col_map and col_map["age"] < len(row_vals) else ""
            employment_type = row_vals[col_map["employment_type"]] if "employment_type" in col_map and col_map["employment_type"] < len(row_vals) else ""
            other_req = row_vals[col_map["other_requirements"]] if "other_requirements" in col_map and col_map["other_requirements"] < len(row_vals) else ""

            # 过滤汇总/统计行
            if re.match(r"^(合计|总计|小计|共计)", job_name.strip()) or (re.match(r"^(合计|总计|小计|共计)", unit.strip()) and not major_raw):
                continue

            if unit:
                last_unit = unit
            else:
                unit = last_unit

            if not job_name and not major_raw:
                continue

            try:
                digits = re.findall(r"\d+", headcount_raw)
                headcount = int(digits[0]) if digits else 1
            except:
                headcount = 1

            full_edu = education
            if degree and degree not in full_edu:
                full_edu = f"{full_edu}（{degree}）" if full_edu else degree

            jobs.append({
                "unit_name": unit or default_unit_name or "未指定招聘单位",
                "job_name": job_name or "未命名岗位",
                "job_code": job_code,
                "headcount": headcount,
                "education": full_edu,
                "major_raw": major_raw,
                "age": age,
                "employment_type": employment_type,
                "other_requirements": other_req
            })
        return jobs
